from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session, send_file
from flask_bcrypt import Bcrypt
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
import os
import json
import uuid
from datetime import datetime
import re
import requests
from bs4 import BeautifulSoup
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx
import docx
import PyPDF2
from youtube_transcript_api import YouTubeTranscriptApi
import urllib.parse
from fpdf import FPDF
from dotenv import load_dotenv
from deep_translator import GoogleTranslator

# Add this import near the top of the file, with the other imports
import chatbot_responses

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'default_secret_key')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///summarization.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size

# Add custom Jinja2 filters
@app.template_filter('regex_search')
def regex_search(text, pattern):
    import re
    return bool(re.search(pattern, text))

@app.template_filter('regex_replace')
def regex_replace(text, pattern, replacement):
    import re
    return re.sub(pattern, replacement, text)

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize extensions
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)

# Download NLTK resources
def download_nltk_resources():
    resources = ['punkt', 'stopwords']
    
    # Add punkt_tab to the list since it's being requested by the code
    if 'punkt_tab' not in resources:
        resources.append('punkt_tab')
        
    for resource in resources:
        try:
            nltk.data.find(f'tokenizers/{resource}')
        except LookupError:
            print(f"Downloading NLTK resource: '{resource}'")
            nltk.download(resource)
            print(f"Resource '{resource}' has been downloaded successfully.")

# Call this function before the app starts
download_nltk_resources()

# Also add this at the top of your app to ensure it's run before any NLTK functionality is used
nltk.download('punkt')
nltk.download('punkt_tab')  # Add this line to download the punkt_tab resource specifically

# DB Models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(60), nullable=False)
    summaries = db.relationship('Summary', backref='author', lazy=True)

class Summary(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    original_text = db.Column(db.Text, nullable=False)
    summary_text = db.Column(db.Text, nullable=False)
    source_type = db.Column(db.String(50), nullable=False)
    compression_ratio = db.Column(db.Float, nullable=False)
    date_created = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

# Helper Functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'txt', 'pdf', 'docx'}

def extract_text_from_file(file):
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    file_type = filename.rsplit('.', 1)[1].lower()
    text = ""
    
    try:
        if file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif file_type == 'pdf':
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text()
        elif file_type == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
    except Exception as e:
        print(f"Error extracting text: {e}")
    finally:
        # Clean up the uploaded file
        if os.path.exists(file_path):
            os.remove(file_path)
    
    return text

def extract_text_from_url(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
            
        # Get text
        text = soup.get_text(separator=' ', strip=True)
        
        # Break into lines and remove leading and trailing space on each
        lines = (line.strip() for line in text.splitlines())
        # Break multi-headlines into a line each
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        # Drop blank lines
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except Exception as e:
        print(f"Error extracting text from URL: {e}")
        return ""

def extract_text_from_youtube(youtube_url):
    try:
        # Extract video ID from URL
        video_id = None
        if 'youtube.com' in youtube_url:
            query = urllib.parse.urlparse(youtube_url).query
            params = dict(urllib.parse.parse_qsl(query))
            video_id = params.get('v')
        elif 'youtu.be' in youtube_url:
            video_id = youtube_url.split('/')[-1].split('?')[0]
        
        if not video_id:
            return "Error: Invalid YouTube URL. Please provide a valid YouTube video URL."
            
        # Get transcript with better error handling
        try:
            transcript = YouTubeTranscriptApi.get_transcript(video_id, languages=['en'])
        except Exception as e:
            # Try to get transcript in any available language if English fails
            try:
                transcript = YouTubeTranscriptApi.get_transcript(video_id)
            except Exception as e:
                if "No transcripts were found" in str(e):
                    return "Error: This video has no captions available. Please try a different video with captions enabled."
                elif "Could not find the requested language" in str(e):
                    return "Error: English captions not available. Please try a video with English captions."
                else:
                    return f"Error: Could not fetch transcript. Please ensure the video has captions enabled and is publicly accessible."
            
        if not transcript:
            return "Error: No transcript content found. Please try a different video."
            
        # Process transcript to combine text and remove timestamps
        text_segments = []
        current_sentence = []
        
        for item in transcript:
            text = item['text'].strip()
            if not text:  # Skip empty segments
                continue
                
            # Add text to current sentence
            current_sentence.append(text)
            
            # Check if the text ends with sentence-ending punctuation
            if text[-1] in '.!?':
                # Join the current sentence and add it to segments
                text_segments.append(' '.join(current_sentence))
                current_sentence = []
            # If sentence is getting too long without punctuation, force a break
            elif len(' '.join(current_sentence).split()) > 20:
                text_segments.append(' '.join(current_sentence) + '.')
                current_sentence = []
        
        # Add any remaining text as a sentence
        if current_sentence:
            text_segments.append(' '.join(current_sentence) + '.')
        
        # Join all segments with proper spacing
        text = ' '.join(text_segments)
        # Clean up multiple spaces and normalize punctuation
        text = re.sub(r'\s+', ' ', text)
        
        if not text.strip():
            return "Error: No readable text could be extracted from the video captions."
            
        return text
        
    except Exception as e:
        return f"Error: An unexpected error occurred while processing the YouTube video. Please try again or use a different video."

def sentence_similarity(sent1, sent2, stopwords=None):
    if stopwords is None:
        stopwords = []
    
    sent1 = [word.lower() for word in sent1]
    sent2 = [word.lower() for word in sent2]
    
    all_words = list(set(sent1 + sent2))
    
    vector1 = [0] * len(all_words)
    vector2 = [0] * len(all_words)
    
    # Build the vector for the first sentence
    for w in sent1:
        if w in stopwords:
            continue
        vector1[all_words.index(w)] += 1
    
    # Build the vector for the second sentence
    for w in sent2:
        if w in stopwords:
            continue
        vector2[all_words.index(w)] += 1
    
    return 1 - cosine_distance(vector1, vector2)

def build_similarity_matrix(sentences, stop_words):
    # Create an empty similarity matrix
    similarity_matrix = np.zeros((len(sentences), len(sentences)))
    
    for idx1 in range(len(sentences)):
        for idx2 in range(len(sentences)):
            if idx1 == idx2:  # ignore if both are same sentences
                continue
            similarity_matrix[idx1][idx2] = sentence_similarity(
                sentences[idx1], sentences[idx2], stop_words)
                
    return similarity_matrix

def generate_summary(text, num_sentences=5):
    stop_words = set(stopwords.words('english'))
    summarize_text = []
    
    # Use custom sentence tokenization for non-Latin scripts
    def tokenize_sentences(text):
        # Check if text contains CJK characters (Chinese, Japanese)
        if any('\u4e00' <= char <= '\u9fff' for char in text) or \
           any('\u3040' <= char <= '\u309f' for char in text) or \
           any('\u30a0' <= char <= '\u30ff' for char in text):
            # Split on common CJK sentence endings
            sentences = []
            current = ""
            for char in text:
                current += char
                if char in ['。', '！', '？', '．', '!', '?', '.']:
                    if current.strip():
                        sentences.append(current.strip())
                    current = ""
            if current.strip():
                sentences.append(current.strip())
            return sentences
        # Check if text contains Hindi characters
        elif any('\u0900' <= char <= '\u097f' for char in text):
            # Split on common Hindi sentence endings
            sentences = []
            current = ""
            for char in text:
                current += char
                if char in ['।', '!', '?', '.']:
                    if current.strip():
                        sentences.append(current.strip())
                    current = ""
            if current.strip():
                sentences.append(current.strip())
            return sentences
        else:
            # Use NLTK for other languages
            return sent_tokenize(text)
    
    # Step 1 - Split text into sentences using custom tokenization
    sentences = tokenize_sentences(text)
    
    # Check if there are enough sentences to summarize
    if len(sentences) <= num_sentences:
        return ' '.join(sentences)
    
    # Step 2 - Tokenize the sentences
    sentence_tokens = [word_tokenize(s.lower()) for s in sentences]
    
    # Step 3 - Create similarity matrix
    sentence_similarity_matrix = build_similarity_matrix(sentence_tokens, stop_words)
    
    # Step 4 - Rank sentences using PageRank algorithm
    sentence_similarity_graph = nx.from_numpy_array(sentence_similarity_matrix)
    scores = nx.pagerank(sentence_similarity_graph)
    
    # Step 5 - Sort the sentences by score and select top n
    ranked_sentences = sorted(((scores[i], s) for i, s in enumerate(sentences)), reverse=True)
    
    for i in range(min(num_sentences, len(ranked_sentences))):
        summarize_text.append(ranked_sentences[i][1])
    
    # Step 6 - Reorder the selected sentences based on their original order
    summarize_text.sort(key=lambda x: sentences.index(x))
    
    # Step 7 - Return the summarized text
    return ' '.join(summarize_text)

def calculate_rouge_score(reference, summary):
    # Simple ROUGE-1 calculation (word overlap)
    reference_words = set(word_tokenize(reference.lower()))
    summary_words = set(word_tokenize(summary.lower()))
    
    overlap = reference_words.intersection(summary_words)
    
    if len(reference_words) == 0:
        return 0
    
    return len(overlap) / len(reference_words)

def format_summary_as_bullets(summary_text):
    # Custom sentence tokenization for different languages
    def tokenize_sentences(text):
        # Check if text contains CJK characters (Chinese, Japanese, Korean)
        if any('\u4e00' <= char <= '\u9fff' for char in text) or \
           any('\u3040' <= char <= '\u309f' for char in text) or \
           any('\u30a0' <= char <= '\u30ff' for char in text) or \
           any('\uac00' <= char <= '\ud7af' for char in text):  # Korean characters
            # Split on common CJK sentence endings
            sentences = []
            current = ""
            for char in text:
                current += char
                if char in ['。', '！', '？', '．', '!', '?', '.', '…', '…', '…']:
                    if current.strip():
                        sentences.append(current.strip())
                    current = ""
            if current.strip():
                sentences.append(current.strip())
            return sentences
        # Check if text contains Hindi characters
        elif any('\u0900' <= char <= '\u097f' for char in text):
            # Split on common Hindi sentence endings
            sentences = []
            current = ""
            for char in text:
                current += char
                if char in ['।', '!', '?', '.', '…']:
                    if current.strip():
                        sentences.append(current.strip())
                    current = ""
            if current.strip():
                sentences.append(current.strip())
            return sentences
        else:
            # Use NLTK for other languages
            return sent_tokenize(text)

    sentences = tokenize_sentences(summary_text)
    
    # Use different bullet points for different languages
    bullet_points = {
        'ja': '・',  # Japanese bullet
        'zh': '•',   # Chinese bullet
        'ko': '•',   # Korean bullet
        'hi': '•',   # Hindi bullet
        'default': '•'  # Default bullet
    }
    
    # Detect language based on character ranges
    def detect_language(text):
        if any('\u4e00' <= char <= '\u9fff' for char in text):  # Chinese characters
            return 'zh'
        elif any('\u3040' <= char <= '\u309f' or '\u30a0' <= char <= '\u30ff' for char in text):  # Japanese characters
            return 'ja'
        elif any('\uac00' <= char <= '\ud7af' for char in text):  # Korean characters
            return 'ko'
        elif any('\u0900' <= char <= '\u097f' for char in text):  # Hindi characters
            return 'hi'
        return 'default'
    
    language = detect_language(summary_text)
    bullet = bullet_points.get(language, bullet_points['default'])
    
    # For CJK and Hindi, ensure each sentence gets its own bullet point
    if language in ['ja', 'zh', 'ko', 'hi']:
        return '\n'.join([f"{bullet} {sentence}" for sentence in sentences])
    else:
        return '\n'.join([f"{bullet} {sentence}" for sentence in sentences])

def format_summary_as_paragraphs(summary_text):
    import re
    def tokenize_sentences(text):
        if any('\u4e00' <= char <= '\u9fff' for char in text) or \
           any('\u3040' <= char <= '\u309f' for char in text) or \
           any('\u30a0' <= char <= '\u30ff' for char in text) or \
           any('\uac00' <= char <= '\ud7af' for char in text):
            sentences = [s.strip() for s in re.split(r'[。！？．!?.]', text) if s.strip()]
            return sentences
        elif any('\u0900' <= char <= '\u097f' for char in text):
            sentences = [s.strip() for s in re.split(r'[।!?\.]', text) if s.strip()]
            return sentences
        else:
            return sent_tokenize(text)

    sentences = tokenize_sentences(summary_text)
    paragraphs = []
    current_paragraph = []
    is_cjk_hi = any('\u4e00' <= char <= '\u9fff' for char in summary_text) or \
                any('\u3040' <= char <= '\u309f' for char in summary_text) or \
                any('\u30a0' <= char <= '\u30ff' for char in summary_text) or \
                any('\uac00' <= char <= '\ud7af' for char in summary_text) or \
                any('\u0900' <= char <= '\u097f' for char in summary_text)
    group_size = 2 if is_cjk_hi else 3
    for i, sentence in enumerate(sentences):
        current_paragraph.append(sentence)
        if (i + 1) % group_size == 0 or i == len(sentences) - 1:
            paragraphs.append(' '.join(current_paragraph))
            current_paragraph = []
    return '\n\n'.join(paragraphs)

def create_docx(title, content, language='en', font_style='Arial', font_size=16):
    import docx
    doc = docx.Document()
    doc.core_properties.language = language

    # Language-specific font mapping
    font_map = {
        'ja': 'MS Gothic',  # Japanese font
        'zh': 'SimSun',     # Chinese font
        'hi': 'Arial Unicode MS',  # Hindi font
        'default': font_style  # Use user-selected font as default
    }

    # Detect language based on content
    def detect_language(text):
        if any('\u4e00' <= char <= '\u9fff' for char in text):  # Chinese characters
            return 'zh'
        elif any('\u3040' <= char <= '\u309f' or '\u30a0' <= char <= '\u30ff' for char in text):  # Japanese characters
            return 'ja'
        elif any('\u0900' <= char <= '\u097f' for char in text):  # Hindi characters
            return 'hi'
        return 'default'

    detected_language = detect_language(content)
    docx_font = font_map.get(detected_language, font_style)

    # Add title
    title_para = doc.add_heading(level=0)
    title_run = title_para.add_run(title)
    try:
        font = title_run.font
        font.name = docx_font
        font.size = docx.shared.Pt(font_size + 4)
    except Exception as e:
        print(f"DOCX Title Font Error: {e}. Using default font.")

    # Process content line by line
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:  # Skip empty lines
            continue
        # Check if line starts with bullet point (including language-specific bullets)
        if line.startswith('•') or line.startswith('・'):
            p = doc.add_paragraph(style='List Bullet')
            text = line[1:].strip()
            run = p.add_run(text)
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
        try:
            font = run.font
            font.name = docx_font
            font.size = docx.shared.Pt(font_size)
        except Exception as e:
            print(f"DOCX Content Font Error: {e}. Using default font.")

    filename = f"summary_{uuid.uuid4()}.docx"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc.save(file_path)
    return file_path

def create_pdf(title, content, language='en', font_style='Arial', font_size=16):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()

    # Expanded font map for PDF
    font_map = {
        'Arial': 'Arial',
        'Helvetica': 'helvetica',
        'Times New Roman': 'Times',
        'Times': 'Times',
        'Courier New': 'Courier',
        'Courier': 'Courier',
        'Verdana': 'Verdana',
        'Georgia': 'Georgia',
        'Palatino': 'Palatino',
        'Garamond': 'Garamond',
        'Bookman': 'Bookman',
        'Comic Sans MS': 'ComicSansMS',
        'Trebuchet MS': 'TrebuchetMS',
        'Arial Black': 'ArialBlack',
        'Impact': 'Impact',
        'Lucida Console': 'LucidaConsole',
        'Tahoma': 'Tahoma',
        'Geneva': 'Geneva',
        'Lucida Sans Unicode': 'LucidaSansUnicode',
        'Noto Sans': 'NotoSans',
        'Noto Sans Kannada': 'NotoSansKannada',
        'Noto Sans CJK': 'NotoSansCJK',
        'Segoe UI': 'SegoeUI',
        'Calibri': 'Calibri',
        'Cambria': 'Cambria',
    }
    fpdf_font = font_map.get(font_style, font_style)

    # Path to font files
    font_dir = os.path.join(os.path.dirname(__file__), 'static', 'fonts')
    font_files = {
        'Verdana': 'Verdana.ttf',
        'Georgia': 'Georgia.ttf',
        'Palatino': 'Palatino.ttf',
        'Garamond': 'Garamond.ttf',
        'Bookman': 'Bookman.ttf',
        'ComicSansMS': 'ComicSansMS.ttf',
        'TrebuchetMS': 'TrebuchetMS.ttf',
        'ArialBlack': 'ArialBlack.ttf',
        'Impact': 'Impact.ttf',
        'LucidaConsole': 'LucidaConsole.ttf',
        'Tahoma': 'Tahoma.ttf',
        'Geneva': 'Geneva.ttf',
        'LucidaSansUnicode': 'LucidaSansUnicode.ttf',
        'NotoSans': 'NotoSans-Regular.ttf',
        'NotoSansKannada': 'NotoSansKannada-Regular.ttf',
        'NotoSansCJK': 'NotoSansCJK-Regular.ttf',
        'SegoeUI': 'SegoeUI.ttf',
        'Calibri': 'calibri.ttf',
        'Cambria': 'cambria.ttf',
    }

    # Register font if custom font is selected and file exists
    try:
        if language in ['zh-CN', 'ja', 'ko']:
            cjk_font_path = os.path.join(font_dir, font_files['NotoSansCJK'])
            if os.path.exists(cjk_font_path):
                pdf.add_font('NotoSansCJK', '', cjk_font_path, uni=True)
                pdf.set_font('NotoSansCJK', size=font_size)
            else:
                pdf.set_font(fpdf_font, size=font_size)
        elif language == 'kn':
            kannada_font_path = os.path.join(font_dir, font_files['NotoSansKannada'])
            if os.path.exists(kannada_font_path):
                pdf.add_font('NotoSansKannada', '', kannada_font_path, uni=True)
                pdf.set_font('NotoSansKannada', size=font_size)
            else:
                pdf.set_font(fpdf_font, size=font_size)
        elif fpdf_font in font_files:
            font_path = os.path.join(font_dir, font_files[fpdf_font])
            if os.path.exists(font_path):
                pdf.add_font(fpdf_font, '', font_path, uni=True)
                pdf.set_font(fpdf_font, size=font_size)
            else:
                pdf.set_font(fpdf_font, size=font_size)
        else:
            pdf.set_font(fpdf_font, size=font_size)
    except Exception as e:
        print(f"PDF Font Error: {e}. Falling back to Arial.")
        pdf.set_font("Arial", size=font_size)

    # Add title
    try:
        pdf.set_font(pdf.font_family, 'B', font_size + 4)
        pdf.cell(200, 10, txt=title, ln=True, align='C')
        pdf.ln(10)
        pdf.set_font(pdf.font_family, '', font_size)
    except Exception as e:
        print(f"PDF Title Font Error: {e}. Falling back to Arial Bold.")
        pdf.set_font("Arial", 'B', font_size + 4)
        pdf.cell(200, 10, txt=title, ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", '', font_size)

    # Add content line by line to preserve font
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        pdf.multi_cell(0, 10, txt=line)

    filename = f"summary_{uuid.uuid4()}.pdf"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    pdf.output(file_path)
    return file_path

def calculate_compression_ratio(original_text, summary_text):
    if not original_text or not summary_text:
        return 0.0  # Return minimum compression ratio instead of 0
    
    # Clean and normalize text before calculating ratio
    original_text = re.sub(r'\s+', ' ', original_text.strip())
    summary_text = re.sub(r'\s+', ' ', summary_text.strip())
    
    # Calculate lengths using word count for more meaningful ratio
    original_words = len(word_tokenize(original_text))
    summary_words = len(word_tokenize(summary_text))
    
    if original_words == 0 or summary_words > original_words:  # Avoid division by zero and invalid ratios
        return 0.0  # Return minimum compression ratio
    
    # Calculate compression (how much text was removed)
    compression = (original_words - summary_words) / original_words
    # Ensure ratio is between 0 and 1
    compression = max(0.0, min(1.0, compression))
    return compression

# Routes
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        valid = False
        if user:
            hash_val = user.password
            if hash_val.startswith("$2b$") or hash_val.startswith("$2a$") or hash_val.startswith("$2y$"):
                try:
                    valid = bcrypt.check_password_hash(hash_val, password)
                except Exception:
                    valid = False
            else:
                # Not a bcrypt hash, treat as invalid or force password reset
                valid = False
        if valid:
            session['user_id'] = user.id
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Login unsuccessful. Please check username and password.', 'danger')
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        
        # Check if username or email already exists
        existing_user = User.query.filter_by(username=username).first()
        existing_email = User.query.filter_by(email=email).first()
        
        if existing_user:
            flash('Username already exists. Please choose a different one.', 'danger')
        elif existing_email:
            flash('Email already exists. Please use a different one.', 'danger')
        else:
            # Hash password and create new user
            hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
            user = User(username=username, email=email, password=hashed_password)
            db.session.add(user)
            db.session.commit()
            flash('Account created successfully! You can now log in.', 'success')
            return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/logout')
def logout():
    session.pop('user_id', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('home'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        flash('Please log in to access the dashboard.', 'warning')
        return redirect(url_for('login'))
        
    user = User.query.get_or_404(session['user_id'])
    summaries = Summary.query.filter_by(user_id=user.id).order_by(Summary.date_created.desc()).all()
    
    return render_template('dashboard.html', user=user, summaries=summaries)

@app.route('/summarize', methods=['GET', 'POST'])
def summarize():
    if 'user_id' not in session:
        flash('Please log in to use the summarization tool.', 'warning')
        return redirect(url_for('login'))
        
    if request.method == 'POST':
        source_type = request.form.get('source_type')
        compression_ratio = float(request.form.get('compression_ratio', 50)) / 100
        view_type = request.form.get('view_type', 'plain')
        
        original_text = ""
        title = "Summary"
        
        if source_type == 'file':
            if 'file' not in request.files:
                flash('No file part', 'danger')
                return redirect(request.url)
                
            file = request.files['file']
            
            if file.filename == '':
                flash('No selected file', 'danger')
                return redirect(request.url)
                
            if file and allowed_file(file.filename):
                original_text = extract_text_from_file(file)
                title = file.filename
            else:
                flash('File type not supported. Please upload a TXT, PDF, or DOCX file.', 'danger')
                return redirect(request.url)
                
        elif source_type == 'url':
            url = request.form.get('url')
            if not url:
                flash('Please enter a URL', 'danger')
                return redirect(request.url)
                
            original_text = extract_text_from_url(url)
            title = url
            
        elif source_type == 'youtube':
            youtube_url = request.form.get('youtube_url')
            if not youtube_url:
                flash('Please enter a YouTube URL', 'danger')
                return redirect(request.url)
                
            try:
                original_text = extract_text_from_youtube(youtube_url)
                if not original_text:
                    flash('Could not extract transcript from the YouTube video. Please ensure the video has captions available.', 'danger')
                    return redirect(request.url)
                title = youtube_url
            except ValueError as e:
                flash(str(e), 'danger')
                return redirect(request.url)
            
        elif source_type == 'text':
            original_text = request.form.get('text', '')
            title = "Manual Text Input"
            
        if not original_text:
            flash('No text could be extracted from the source.', 'danger')
            return redirect(request.url)
            
        # Calculate number of sentences based on compression ratio
        sentences = sent_tokenize(original_text)
        num_sentences = max(1, int(len(sentences) * (1 - compression_ratio)))
        
        # Generate summary
        summary_text = generate_summary(original_text, num_sentences)
        
        # Calculate actual compression ratio based on word count
        actual_compression = calculate_compression_ratio(original_text, summary_text)
        
        # Format summary based on view type
        formatted_summary = summary_text
        if view_type == 'bullet':
            formatted_summary = format_summary_as_bullets(summary_text)
        elif view_type == 'paragraph':
            formatted_summary = format_summary_as_paragraphs(summary_text)
        
        # Calculate metrics
        word_count = len(word_tokenize(summary_text))
        sentence_count = len(sent_tokenize(summary_text))
        rouge_score = calculate_rouge_score(original_text, summary_text)
        
        # Save summary to database
        new_summary = Summary(
            title=title,
            original_text=original_text,
            summary_text=summary_text,
            source_type=source_type,
            compression_ratio=actual_compression,
            user_id=session['user_id']
        )
        db.session.add(new_summary)
        db.session.commit()
        
        metrics = {
            'word_count': word_count,
            'sentence_count': sentence_count,
            'compression_ratio': actual_compression,
            'rouge_score': rouge_score
        }
        
        return render_template(
            'summary_result.html',
            title=title,
            original_text=original_text,
            summary_text=formatted_summary,
            metrics=metrics,
            view_type=view_type,
            summary_id=new_summary.id
        )
        
    return render_template('summarize.html')

@app.route('/summary/<int:summary_id>')
def view_summary(summary_id):
    if 'user_id' not in session:
        flash('Please log in to view summaries.', 'warning')
        return redirect(url_for('login'))
        
    summary = Summary.query.get_or_404(summary_id)
    
    # Check if the summary belongs to the logged-in user
    if summary.user_id != session['user_id']:
        flash('You do not have permission to view this summary.', 'danger')
        return redirect(url_for('dashboard'))
    
    # Calculate metrics
    word_count = len(word_tokenize(summary.summary_text))
    sentence_count = len(sent_tokenize(summary.summary_text))
    rouge_score = calculate_rouge_score(summary.original_text, summary.summary_text)
    
    metrics = {
        'word_count': word_count,
        'sentence_count': sentence_count,
        'compression_ratio': summary.compression_ratio,
        'rouge_score': rouge_score
    }
    
    return render_template(
        'summary_result.html',
        title=summary.title,
        original_text=summary.original_text,
        summary_text=summary.summary_text,
        metrics=metrics,
        view_type='plain',
        summary_id=summary.id
    )

@app.route('/format_summary', methods=['POST'])
def format_summary():
    import re
    summary_text = request.form.get('summary_text')
    view_type = request.form.get('view_type')

    def detect_language(text):
        if any('\u4e00' <= char <= '\u9fff' for char in text):  # Chinese
            return 'zh'
        elif any('\u3040' <= char <= '\u309f' or '\u30a0' <= char <= '\u30ff' for char in text):  # Japanese
            return 'ja'
        elif any('\u0900' <= char <= '\u097f' for char in text):  # Hindi
            return 'hi'
        return 'en'

    language = detect_language(summary_text)

    def custom_tokenize(text, lang):
        if lang == 'ja' or lang == 'zh':
            sentences = [s.strip() for s in re.split(r'[。！？．!?。]', text) if s.strip()]
            if len(sentences) <= 1 and len(text) > 40:
                sentences = [text[i:i+40] for i in range(0, len(text), 40)]
            return sentences
        elif lang == 'hi':
            sentences = [s.strip() for s in re.split(r'[।!?\.]', text) if s.strip()]
            if len(sentences) <= 1 and len(text) > 30:
                sentences = [text[i:i+30] for i in range(0, len(text), 30)]
            return sentences
        else:
            return None  # fallback to default

    if language in ['ja', 'zh', 'hi']:
        sentences = custom_tokenize(summary_text, language)
        if view_type == 'bullet':
            bullet = '・' if language == 'ja' else '•'
            formatted_text = '\n'.join([f"{bullet} {sentence}" for sentence in sentences])
        elif view_type == 'paragraph':
            paragraphs = []
            current_paragraph = []
            for i, sentence in enumerate(sentences):
                current_paragraph.append(sentence)
                if (i + 1) % 3 == 0 or i == len(sentences) - 1:
                    paragraphs.append(' '.join(current_paragraph))
                    current_paragraph = []
            formatted_text = '\n\n'.join(paragraphs)
        else:
            formatted_text = summary_text
    else:
        # Use default helpers for other languages
        if view_type == 'bullet':
            formatted_text = format_summary_as_bullets(summary_text)
        elif view_type == 'paragraph':
            formatted_text = format_summary_as_paragraphs(summary_text)
        else:
            formatted_text = summary_text

    return jsonify({'formatted_text': formatted_text})

@app.route('/export_summary/<int:summary_id>', methods=['POST'])
def export_summary(summary_id):
    if 'user_id' not in session:
        flash('Please log in to export summaries.', 'warning')
        return redirect(url_for('login'))
        
    summary = Summary.query.get_or_404(summary_id)
    
    # Check if the summary belongs to the logged-in user
    if summary.user_id != session['user_id']:
        flash('You do not have permission to export this summary.', 'danger')
        return redirect(url_for('dashboard'))
    
    export_format = request.form.get('export_format')
    target_language = request.form.get('target_language', 'en')
    font_size = request.form.get('font_size', '16px')
    font_style = request.form.get('font_style', 'Arial')
    view_type = request.form.get('view_type', 'plain')
    file_path = None
    
    try:
        # Get the current translated text if available, otherwise use original summary
        current_text = request.form.get('current_text', summary.summary_text)
        
        # Detect language and format text accordingly
        def detect_language(text):
            if any('\u4e00' <= char <= '\u9fff' for char in text):  # Chinese
                return 'zh'
            elif any('\u3040' <= char <= '\u309f' or '\u30a0' <= char <= '\u30ff' for char in text):  # Japanese
                return 'ja'
            elif any('\u0900' <= char <= '\u097f' for char in text):  # Hindi
                return 'hi'
            return 'en'
        
        language = detect_language(current_text)
        
        # Format the text based on view_type and language
        if view_type == 'bullet':
            if language in ['ja', 'zh', 'hi']:
                # Custom bullet formatting for CJK and Hindi
                sentences = []
                current = ""
                for char in current_text:
                    current += char
                    if (language in ['ja', 'zh'] and char in ['。', '！', '？', '．', '!', '?', '.']) or \
                       (language == 'hi' and char in ['।', '!', '?', '.']):
                        if current.strip():
                            sentences.append(current.strip())
                        current = ""
                if current.strip():
                    sentences.append(current.strip())
                
                bullet = '・' if language == 'ja' else '•'
                formatted_text = '\n'.join([f"{bullet} {sentence}" for sentence in sentences])
            else:
                formatted_text = format_summary_as_bullets(current_text)
                
        elif view_type == 'paragraph':
            if language in ['ja', 'zh', 'hi']:
                # Custom paragraph formatting for CJK and Hindi
                sentences = []
                current = ""
                for char in current_text:
                    current += char
                    if (language in ['ja', 'zh'] and char in ['。', '！', '？', '．', '!', '?', '.']) or \
                       (language == 'hi' and char in ['।', '!', '?', '.']):
                        if current.strip():
                            sentences.append(current.strip())
                        current = ""
                if current.strip():
                    sentences.append(current.strip())
                
                # Group sentences into paragraphs
                paragraphs = []
                current_paragraph = []
                for i, sentence in enumerate(sentences):
                    current_paragraph.append(sentence)
                    if (i + 1) % 3 == 0 or i == len(sentences) - 1:
                        paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
                
                formatted_text = '\n\n'.join(paragraphs)
            else:
                formatted_text = format_summary_as_paragraphs(current_text)
        else:  # plain
            formatted_text = current_text
        
        # Convert font size from px to int (default 16)
        try:
            font_size_int = int(font_size.replace('px', ''))
        except Exception:
            font_size_int = 16
        
        if export_format == 'txt':
            # Create a text file with font info as header
            filename = f"summary_{uuid.uuid4()}.txt"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"Font: {font_style}, Size: {font_size}\n\n")
                f.write(f"{summary.title}\n\n")
                f.write(formatted_text)
        elif export_format == 'pdf':
            try:
                # Always create DOCX first
                docx_path = create_docx(summary.title, formatted_text, target_language, font_style, font_size_int)
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f"summary_{uuid.uuid4()}.pdf")
                conversion_success = False

                # Try docx2pdf
                try:
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                    conversion_success = os.path.exists(pdf_path)
                except Exception as e:
                    print(f"docx2pdf conversion failed: {e}")

                # If docx2pdf failed, try using comtypes (Windows only)
                if not conversion_success:
                    try:
                        import comtypes.client
                        import pythoncom
                        pythoncom.CoInitialize()
                        word = comtypes.client.CreateObject('Word.Application')
                        doc = word.Documents.Open(os.path.abspath(docx_path))
                        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
                        doc.Close()
                        word.Quit()
                        conversion_success = os.path.exists(pdf_path)
                    except Exception as e:
                        print(f"comtypes conversion failed: {e}")

                # (Optional) Try using libreoffice as a last resort (Linux)
                if not conversion_success:
                    try:
                        import subprocess
                        subprocess.run([
                            'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                            os.path.dirname(pdf_path), docx_path
                        ], check=True)
                        # LibreOffice names the PDF the same as the DOCX but with .pdf
                        libre_pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
                        if os.path.exists(libre_pdf_path):
                            os.rename(libre_pdf_path, pdf_path)
                            conversion_success = True
                    except Exception as e:
                        print(f"libreoffice conversion failed: {e}")

                # If all conversions failed, fall back to direct PDF (may lose formatting)
                if not conversion_success:
                    print('All DOCX to PDF conversions failed, falling back to direct PDF.')
                    pdf_path = create_pdf(summary.title, formatted_text, target_language, font_style, font_size_int)

                # Clean up the temporary DOCX file
                if os.path.exists(docx_path):
                    os.remove(docx_path)

                file_path = pdf_path

            except Exception as e:
                print(f"Error in PDF conversion: {str(e)}")
                flash('Error creating PDF. Please try again.', 'danger')
                return redirect(url_for('view_summary', summary_id=summary_id))
        elif export_format == 'docx':
            file_path = create_docx(summary.title, formatted_text, target_language, font_style, font_size_int)
        
        if file_path:
            return send_file(file_path, as_attachment=True)
        else:
            flash('Error exporting summary.', 'danger')
            return redirect(url_for('view_summary', summary_id=summary_id))
            
    except Exception as e:
        print(f"Export error: {str(e)}")
        flash('Error exporting summary. Please try again.', 'danger')
        return redirect(url_for('view_summary', summary_id=summary_id))

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    if request.method == 'POST':
        # Process contact form submission (would typically send an email)
        flash('Thank you for your message! We will respond shortly.', 'success')
        return redirect(url_for('contact'))
    return render_template('contact.html')

@app.route('/privacy')
def privacy():
    return render_template('privacy.html')

@app.route('/terms')
def terms():
    return render_template('terms.html')

@app.route('/pricing')
def pricing():
    return render_template('pricing.html')

@app.route('/faq')
def faq():
    return render_template('faq.html')

@app.route('/cookies')
def cookies():
    return render_template('cookies.html')

@app.route('/gdpr')
def gdpr():
    return render_template('gdpr.html')


from flask import send_from_directory

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'), 'img/favicon.ico', mimetype='image/vnd.microsoft.icon')

# Replace the entire chatbot route with this simpler version
@app.route('/chatbot', methods=['POST'])
def chatbot():
    data = request.get_json()
    user_message = data.get('message', '')
    
    # Get response from the chatbot_responses module
    response = chatbot_responses.get_response(user_message)
    
    return jsonify({'response': response})

@app.route('/translate_summary', methods=['POST'])
def translate_summary():
    if 'user_id' not in session:
        return jsonify({'error': 'Please log in to translate summaries.'}), 401
        
    data = request.get_json()
    text = data.get('text', '')
    target_language = data.get('target_language', '')
    source_language = data.get('source_language', 'auto')
    
    if not text or not target_language:
        return jsonify({'error': 'Missing required parameters'}), 400
        
    try:
        # First try with explicit source language
        translator = GoogleTranslator(source=source_language, target=target_language)
        
        # Split text into smaller chunks if it's too long (API limit is 5000 chars)
        max_chunk_size = 4000
        text_chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        
        # Translate each chunk and combine
        translated_chunks = []
        for chunk in text_chunks:
            try:
                translated_chunk = translator.translate(chunk)
                translated_chunks.append(translated_chunk)
            except Exception as chunk_error:
                print(f"Chunk translation error with source {source_language}: {str(chunk_error)}")
                # If chunk translation fails, try with auto source language
                translator = GoogleTranslator(source='auto', target=target_language)
                translated_chunk = translator.translate(chunk)
                translated_chunks.append(translated_chunk)
        
        translated_text = ' '.join(translated_chunks)
        
        # Verify the translation was successful
        if not translated_text or translated_text.isspace():
            raise Exception("Translation resulted in empty text")
            
        return jsonify({
            'translated_text': translated_text,
            'source_language': source_language,
            'target_language': target_language
        })
            
    except Exception as e:
        print(f"Translation error: {str(e)}")
        # Try one more time with auto source language
        try:
            translator = GoogleTranslator(source='auto', target=target_language)
            translated_text = translator.translate(text)
            
            if not translated_text or translated_text.isspace():
                raise Exception("Translation resulted in empty text")
                
            return jsonify({
                'translated_text': translated_text,
                'source_language': 'auto',
                'target_language': target_language
            })
        except Exception as retry_error:
            print(f"Retry translation error: {str(retry_error)}")
            return jsonify({'error': 'Translation failed. Please try again or reset to original text.'}), 500

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email')
        user = User.query.filter_by(email=email).first()
        if user:
            flash('A password reset link has been sent to your email (demo: not actually sent).', 'info')
        else:
            flash('No account found with that email address.', 'danger')
    return render_template('forgot_password.html')

# Create tables before first request
@app.before_request
def create_tables_once():
    if not hasattr(app, 'tables_created'):
        app.tables_created = True
        db.create_all()

if __name__ == '__main__':
    app.run(debug=True)
